import requests
from bs4 import BeautifulSoup
import openai
import docx
from sympy import latex, sympify
from docx.shared import Inches
from io import BytesIO
from PIL import Image
import matplotlib.pyplot as plt

def write_to_word(filename, content):
    """Create a Word file with Wikipedia content, headings, text, and LaTeX equations."""
    doc = docx.Document()

    for data in content:
        query, definition, headings, subheadings, formulas, summary = data
        
        # Add topic as a heading
        doc.add_heading(f'Topic: {query}', level=1)
        
        # Add definition
        doc.add_heading('Definition:', level=2)
        doc.add_paragraph(definition)
        # Add Headings and their explanations
        if headings:
            doc.add_heading('Headings:', level=2)
            for heading, explanation in headings.items():
                doc.add_heading(heading, level=3)
                doc.add_paragraph(explanation)
        
        # Add Subheadings and their explanations
        if subheadings:
            doc.add_heading('Subheadings:', level=2)
            for subheading, explanation in subheadings.items():
                doc.add_heading(subheading, level=3)
                doc.add_paragraph(explanation)
        
        # Add Formulas and convert them to LaTeX
        if formulas:
            doc.add_heading('Formulas:', level=2)
            for formula in formulas:
                doc.add_paragraph(f"Formula (LaTeX): {formula}")
                # Convert formula to LaTeX and insert image into Word document
                insert_formula_as_image(doc, formula)

        # Add Summary (first 5 paragraphs)
        doc.add_heading('Summary:', level=2)
        for point in summary:
            doc.add_paragraph(point)

    doc.save(filename)
    print(f"Saved Word document: {filename}")

def create_pdf(filename, content):
    """Create a PDF file with Wikipedia content, headings, text, and LaTeX equations."""
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    text = c.beginText(1 * inch, height - 1 * inch)
    text.setFont("Helvetica", 10)

    for data in content:
        query, definition, headings, subheadings, formulas, summary = data
        
        # Add topic as a heading
        text.setFont("Helvetica-Bold", 14)
        text.textLine(f'Topic: {query}')
        
        # Add definition
        text.setFont("Helvetica-Bold", 12)
        text.textLine('Definition:')
        text.setFont("Helvetica", 10)
        text.textLines(definition)

        # Add Headings and their explanations
        if headings:
            text.setFont("Helvetica-Bold", 12)
            text.textLine('Headings:')
            text.setFont("Helvetica", 10)
            for heading, explanation in headings.items():
                text.setFont("Helvetica-Bold", 11)
                text.textLine(heading)
                text.setFont("Helvetica", 10)
                text.textLines(explanation)
        
        # Add Subheadings and their explanations
        if subheadings:
            text.setFont("Helvetica-Bold", 12)
            text.textLine('Subheadings:')
            text.setFont("Helvetica", 10)
            for subheading, explanation in subheadings.items():
                text.setFont("Helvetica-Bold", 11)
                text.textLine(subheading)
                text.setFont("Helvetica", 10)
                text.textLines(explanation)

        # Add Formulas and convert them to LaTeX
        if formulas:
            text.setFont("Helvetica-Bold", 12)
            text.textLine('Formulas:')
            text.setFont("Helvetica", 10)
            for formula in formulas:
                text.textLine(f"Formula (LaTeX): {formula}")
                # Convert formula to LaTeX and insert image into PDF
                image_stream = insert_formula_as_image(formula)
                if image_stream:
                    c.drawImage(image_stream, 1 * inch, height - (text.getY() + 0.5 * inch), width=2 * inch, height=0.5 * inch)
                    text.moveCursor(0, 50)

        # Add Summary (first 5 paragraphs)
        text.setFont("Helvetica-Bold", 12)
        text.textLine('Summary:')
        text.setFont("Helvetica", 10)
        for point in summary:
            text.textLines(point)

        text.moveCursor(0, 20)
        c.drawText(text)
        c.showPage()
        text = c.beginText(1 * inch, height - 1 * inch)  # Reset text for the next page

    c.save()
    print(f"Saved PDF document: {filename}")

def create_ppt(filename, content):
    """Create a PowerPoint presentation with Wikipedia content, headings, text, and LaTeX equations."""
    prs = Presentation()

    for data in content:
        query, definition, headings, subheadings, formulas, summary = data

        # Create a title slide
        slide = prs.slides.add_slide(prs.slide_layouts[0])
        title = slide.shapes.title
        subtitle = slide.placeholders[1]
        title.text = query
        subtitle.text = definition

        # Add a slide for headings
        slide = prs.slides.add_slide(prs.slide_layouts[1])
        title = slide.shapes.title
        title.text = "Headings"
        content_box = slide.shapes.placeholders[1].text_frame

        for heading, explanation in headings.items():
            p = content_box.add_paragraph()
            p.text = f"{heading}: {explanation}"
            p.space_after = Pt(14)

        # Add a slide for subheadings
        slide = prs.slides.add_slide(prs.slide_layouts[1])
        title = slide.shapes.title
        title.text = "Subheadings"
        content_box = slide.shapes.placeholders[1].text_frame

        for subheading, explanation in subheadings.items():
            p = content_box.add_paragraph()
            p.text = f"{subheading}: {explanation}"
            p.space_after = Pt(14)

        # Add a slide for formulas
        if formulas:
            slide = prs.slides.add_slide(prs.slide_layouts[1])
            title = slide.shapes.title
            title.text = "Formulas"
            content_box = slide.shapes.placeholders[1].text_frame

            for formula in formulas:
                p = content_box.add_paragraph()
                p.text = f"Formula (LaTeX): {formula}"
                img_stream = insert_formula_as_image(formula)
                if img_stream:
                    img = slide.shapes.add_picture(img_stream, Inches(1), Inches(2), width=Inches(8), height=Inches(2))

        # Add a slide for summary
        slide = prs.slides.add_slide(prs.slide_layouts[1])
        title = slide.shapes.title
        title.text = "Summary"
        content_box = slide.shapes.placeholders[1].text_frame

        for point in summary:
            p = content_box.add_paragraph()
            p.text = point
            p.space_after = Pt(14)

    prs.save(filename)
    print(f"Saved PowerPoint presentation: {filename}")

def insert_formula_as_image(doc, formula):
    """Render LaTeX formula as an image and insert it into the Word document."""
    # Use sympy to render LaTeX
    expr = sympify(formula.replace('$', '').replace('\\cdot ', '*'))
    latex_code = latex(expr)

    # Create a plot image for the LaTeX formula
    fig, ax = plt.subplots()
    ax.text(0.5, 0.5, f"${latex_code}$", fontsize=20, ha='center')
    ax.axis('off')
    
    # Save the figure to a BytesIO object
    img_stream = BytesIO()
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)
    plt.close(fig)

    # Insert image into the Word document
    image = Image.open(img_stream)
    image.save(img_stream, format="PNG")
    doc.add_picture(img_stream, width=Inches(3))  # Insert picture into the doc
